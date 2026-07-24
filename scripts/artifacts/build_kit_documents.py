from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "assets" / "kit-templates"
LOGO = ROOT / "app" / "icon.png"

BLUE = "1571D4"
NAVY = "16263B"
ORANGE = "FD8417"
PALE_BLUE = "EAF3FC"
PALE_ORANGE = "FFF1E4"
GRAY = "5F6B7A"
LIGHT_GRAY = "D9E2EC"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(round(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(round(sum(widths) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, width in enumerate(widths):
            set_cell_width(row.cells[index], width)
            set_cell_margins(row.cells[index])
            row.cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_bottom_border(paragraph, color: str = LIGHT_GRAY, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_document(title: str, subtitle: str, audience: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in (
        ("Title", 27, NAVY, 0, 8),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.97))
    set_table_widths(header_table, [1.05, 5.92])
    header_table.cell(0, 0).paragraphs[0].add_run().add_picture(
        str(LOGO), width=Inches(0.48)
    )
    brand = header_table.cell(0, 1).paragraphs[0]
    brand.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = brand.add_run("FORMÉDUCWEB  ·  KIT D’EXÉCUTION 90 JOURS")
    set_run_font(run, size=8, color=GRAY, bold=True)
    add_bottom_border(brand)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(
        "ForméducWeb  ·  Gabarit opérationnel  ·  À adapter et faire valider selon votre contexte"
    )
    set_run_font(run, size=8, color=GRAY)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("KIT D’EXÉCUTION 90 JOURS")
    set_run_font(run, size=9, color=ORANGE, bold=True)

    title_p = doc.add_paragraph(style="Title")
    title_p.add_run(title)

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_after = Pt(12)
    run = subtitle_p.add_run(subtitle)
    set_run_font(run, size=12.5, color=GRAY)

    metadata = doc.add_table(rows=2, cols=2)
    set_table_widths(metadata, [3.485, 3.485])
    metadata.style = "Table Grid"
    values = [
        ("Préparé pour", "{{COMPANY_NAME}}"),
        ("Responsable", "{{CONTACT_NAME}}"),
        ("Date", "{{GENERATED_DATE}}"),
        ("Public visé", audience),
    ]
    for idx, (label, value) in enumerate(values):
        cell = metadata.cell(idx // 2, idx % 2)
        cell.text = ""
        label_p = cell.paragraphs[0]
        label_p.paragraph_format.space_after = Pt(1)
        label_run = label_p.add_run(label.upper())
        set_run_font(label_run, size=7.5, color=GRAY, bold=True)
        value_p = cell.add_paragraph()
        value_p.paragraph_format.space_after = Pt(0)
        value_run = value_p.add_run(value)
        set_run_font(value_run, size=10, color=NAVY, bold=True)
        set_cell_shading(cell, PALE_BLUE if idx % 2 == 0 else "F8FAFC")

    add_callout(
        doc,
        "Priorité issue de votre diagnostic",
        "{{TOP_ACTION_1}}",
        PALE_BLUE,
    )

    return doc


def add_callout(doc: Document, title: str, body: str, fill: str = PALE_ORANGE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [6.97])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    set_run_font(run, size=10, color=NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run = p2.add_run(body)
    set_run_font(run, size=9.5, color=GRAY)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_signoff(doc: Document) -> None:
    doc.add_heading("Adoption et révision", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [1.85, 5.12])
    for row, values in enumerate(
        [
            ("Propriétaire", "{{CONTACT_NAME}}"),
            ("Approuvé par", "À compléter"),
            ("Date d’adoption", "{{GENERATED_DATE}}"),
            ("Prochaine révision", "Dans 90 jours, puis annuellement"),
        ]
    ):
        table.cell(row, 0).text = values[0]
        table.cell(row, 1).text = values[1]
        set_cell_shading(table.cell(row, 0), PALE_BLUE)
        for run in table.cell(row, 0).paragraphs[0].runs:
            set_run_font(run, bold=True, color=NAVY)


def build_ai_charter() -> None:
    doc = configure_document(
        "Charte d’utilisation responsable de l’IA",
        "Des règles simples, applicables dès maintenant, pour utiliser l’IA avec discernement.",
        "Personnel, gestionnaires et fournisseurs autorisés",
    )
    add_callout(
        doc,
        "Mode d’emploi",
        "Remplacez les mentions « À compléter », confirmez les outils autorisés et faites approuver la charte avant diffusion.",
    )
    doc.add_heading("1. Objet et principes", level=1)
    doc.add_paragraph(
        "{{COMPANY_NAME}} utilise l’intelligence artificielle comme outil d’assistance. "
        "Une personne demeure responsable de chaque décision, publication et livraison."
    )
    add_bullets(
        doc,
        [
            "Utilité : partir d’un besoin concret et mesurable.",
            "Minimisation : transmettre le moins de données possible.",
            "Validation humaine : vérifier faits, sources, calculs et ton.",
            "Transparence : signaler l’usage de l’IA lorsque le contexte l’exige.",
            "Traçabilité : consigner les usages significatifs dans le registre prévu.",
        ],
    )
    doc.add_heading("2. Usages permis", level=1)
    add_bullets(
        doc,
        [
            "Idéation, plans, résumés et premières versions de contenus non sensibles.",
            "Aide à la structuration, à la traduction et à l’analyse de documents approuvés.",
            "Automatisation encadrée d’une tâche réversible, testée avant mise en service.",
        ],
    )
    doc.add_heading("3. Données interdites sans autorisation", level=1)
    add_bullets(
        doc,
        [
            "Renseignements personnels ou données permettant d’identifier une personne.",
            "Secrets commerciaux, renseignements financiers, juridiques ou RH.",
            "Mots de passe, clés d’accès, configurations de sécurité ou données clients.",
            "Œuvres ou documents dont les droits d’utilisation ne sont pas confirmés.",
        ],
    )
    doc.add_heading("4. Contrôle avant utilisation", level=1)
    add_numbered(
        doc,
        [
            "Confirmer que l’outil figure dans la liste des outils approuvés.",
            "Retirer ou anonymiser les données non nécessaires.",
            "Vérifier la sortie avec une source fiable ou une personne compétente.",
            "Consigner l’usage s’il influence un client, une décision ou un processus.",
            "Escalader tout doute au responsable désigné.",
        ],
    )
    doc.add_heading("5. Outils et responsabilités", level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    set_table_widths(table, [1.65, 2.75, 2.57])
    headers = ["Élément", "Décision", "Responsable"]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
        set_cell_shading(table.cell(0, idx), BLUE)
        for run in table.cell(0, idx).paragraphs[0].runs:
            set_run_font(run, color=WHITE, bold=True)
    rows = [
        ("Outils approuvés", "À compléter", "{{CONTACT_NAME}}"),
        ("Revue des usages", "Mensuelle pendant 90 jours", "{{CONTACT_NAME}}"),
        ("Incident ou erreur", "Suspendre, documenter, corriger", "Gestionnaire"),
    ]
    for row_idx, row in enumerate(rows, 1):
        for col_idx, value in enumerate(row):
            table.cell(row_idx, col_idx).text = value
    add_signoff(doc)
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT / "charte-ia-editable.docx")


def build_ai_memo() -> None:
    doc = configure_document(
        "Mémo équipe · 7 réflexes IA",
        "Une page de rappel à partager après l’adoption de la charte.",
        "Toute l’équipe",
    )
    doc.add_heading("Avant d’utiliser un outil IA", level=1)
    add_numbered(
        doc,
        [
            "Je pars d’un objectif clair, pas d’un outil à la mode.",
            "Je n’entre aucune donnée personnelle, client ou confidentielle sans autorisation.",
            "J’utilise seulement un outil approuvé par {{COMPANY_NAME}}.",
            "Je vérifie les faits, les chiffres, les sources et les droits d’utilisation.",
            "Je garde une personne responsable de la décision finale.",
            "Je consigne l’usage s’il touche un client, un processus ou une décision.",
            "Je signale rapidement toute erreur, fuite possible ou résultat préoccupant.",
        ],
    )
    add_callout(
        doc,
        "Le bon réflexe",
        "En cas de doute, retirez les données, arrêtez l’envoi et demandez une validation à {{CONTACT_NAME}}.",
        PALE_BLUE,
    )
    doc.add_heading("Mini-test avant publication", level=1)
    checklist = [
        "La sortie répond-elle vraiment au besoin?",
        "Puis-je expliquer comment elle a été validée?",
        "Une donnée sensible pourrait-elle être déduite?",
        "Le ton, les faits et les droits sont-ils acceptables?",
    ]
    for item in checklist:
        doc.add_paragraph(f"☐ {item}")
    doc.add_paragraph(
        "Ce mémo ne remplace pas la charte complète ni une validation juridique, "
        "de sécurité ou de confidentialité adaptée au contexte."
    )
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT / "memo-equipe-ia.docx")


def build_cyber_incident() -> None:
    doc = configure_document(
        "Procédure de gestion d’un incident cyber",
        "Un parcours court pour signaler, contenir, documenter et reprendre les opérations.",
        "Personnel, gestionnaires et fournisseur TI autorisé",
    )
    add_callout(
        doc,
        "Déclencheur",
        "Appliquez cette procédure dès qu’un compte, appareil, paiement, fichier ou service semble compromis. Ne cherchez pas à tout résoudre seul.",
    )
    doc.add_heading("Canal de signalement", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [1.85, 5.12])
    for row, values in enumerate(
        [
            ("Responsable", "{{CONTACT_NAME}}"),
            ("Canal principal", "À compléter · ne pas utiliser un canal compromis"),
            ("Fournisseur TI", "À compléter"),
        ]
    ):
        table.cell(row, 0).text = values[0]
        table.cell(row, 1).text = values[1]
        set_cell_shading(table.cell(row, 0), PALE_BLUE)
    doc.add_heading("Étapes", level=1)
    add_numbered(
        doc,
        [
            "SIGNALER — noter l’heure, le système, l’utilisateur et ce qui a été observé.",
            "CONTENIR — isoler l’appareil ou suspendre l’accès selon les instructions; préserver les preuves.",
            "QUALIFIER — déterminer les services, données et personnes potentiellement touchés.",
            "ESCALADER — mobiliser la direction, le fournisseur TI et les responsables concernés.",
            "CORRIGER — supprimer la cause, réinitialiser les accès et appliquer les correctifs validés.",
            "REPRENDRE — restaurer progressivement, tester et surveiller les services essentiels.",
            "APPRENDRE — consigner la chronologie, les décisions et les mesures préventives.",
        ],
    )
    doc.add_page_break()
    doc.add_heading("Fiche de première heure", level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [2.05, 4.92])
    for row, label in enumerate(
        [
            "Date et heure",
            "Personne ayant signalé",
            "Système ou compte touché",
            "Symptômes observés",
            "Actions déjà prises",
            "Décideur de reprise",
        ]
    ):
        table.cell(row, 0).text = label
        table.cell(row, 1).text = "À compléter"
        set_cell_shading(table.cell(row, 0), PALE_BLUE)
    add_callout(
        doc,
        "À éviter",
        "Ne pas effacer les journaux, répondre à l’attaquant, payer, restaurer ou communiquer publiquement sans décision autorisée.",
        PALE_ORANGE,
    )
    doc.add_heading("Retour d’expérience", level=1)
    add_bullets(
        doc,
        [
            "Cause probable et contrôles qui ont fonctionné.",
            "Temps de détection, de confinement et de reprise.",
            "Actions correctives, responsable et échéance.",
            "Mise à jour de la checklist accès/sauvegardes.",
        ],
    )
    add_signoff(doc)
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT / "procedure-incident-cyber.docx")


def build_cyber_memo() -> None:
    doc = configure_document(
        "Mémo antifraude · ralentir, vérifier, confirmer",
        "Une page de réflexes concrets contre l’hameçonnage et la fraude au paiement.",
        "Toute personne qui reçoit des demandes, factures ou changements bancaires",
    )
    doc.add_heading("Les signaux d’alerte", level=1)
    add_bullets(
        doc,
        [
            "Urgence inhabituelle, secret demandé ou pression pour contourner une règle.",
            "Changement de compte bancaire, de bénéficiaire ou de coordonnées.",
            "Adresse proche de l’originale, pièce jointe inattendue ou ton inhabituel.",
            "Demande de code MFA, mot de passe, carte-cadeau ou cryptomonnaie.",
        ],
    )
    doc.add_heading("La règle des deux canaux", level=1)
    add_numbered(
        doc,
        [
            "Arrêter l’action et ne pas répondre dans le fil suspect.",
            "Retrouver soi-même un numéro ou canal déjà connu.",
            "Confirmer la demande avec une personne autorisée.",
            "Documenter la validation avant tout paiement ou changement.",
        ],
    )
    add_callout(
        doc,
        "Aucun blâme pour un signalement prudent",
        "{{COMPANY_NAME}} préfère un délai de validation à une fraude. Signalez même si vous avez déjà cliqué ou répondu.",
        PALE_BLUE,
    )
    doc.add_heading("Si vous avez cliqué ou transmis une information", level=1)
    add_bullets(
        doc,
        [
            "Déconnectez l’appareil du réseau seulement si la procédure l’exige.",
            "Contactez immédiatement {{CONTACT_NAME}} par le canal prévu.",
            "Ne supprimez pas le message; notez l’heure et les actions effectuées.",
            "N’effectuez pas vous-même de réinitialisation non coordonnée.",
        ],
    )
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT / "memo-antifraude.docx")


def build_law25_requests() -> None:
    doc = configure_document(
        "Procédure de demandes de renseignements personnels",
        "Recevoir, authentifier, traiter et documenter une demande d’accès, de rectification ou de retrait.",
        "Responsable de la protection des renseignements personnels et personnes désignées",
    )
    add_callout(
        doc,
        "Limite importante",
        "Ce gabarit facilite l’organisation interne. Il ne remplace pas un avis juridique ni l’analyse des délais et exceptions applicables.",
    )
    doc.add_heading("1. Réception et journalisation", level=1)
    add_numbered(
        doc,
        [
            "Centraliser la demande au canal officiel de {{COMPANY_NAME}}.",
            "Consigner la date, le type de demande et une référence interne.",
            "Accuser réception sans confirmer de renseignement sensible.",
            "Identifier le responsable du traitement et l’échéance applicable.",
        ],
    )
    doc.add_heading("2. Vérification de l’identité", level=1)
    doc.add_paragraph(
        "Demander uniquement l’information raisonnablement nécessaire pour vérifier l’identité. "
        "Éviter de recueillir une copie complète d’un document si une méthode moins intrusive suffit."
    )
    doc.add_heading("3. Recherche et décision", level=1)
    add_numbered(
        doc,
        [
            "Repérer les systèmes, fournisseurs et équipes susceptibles de détenir l’information.",
            "Préserver les données pertinentes pendant le traitement.",
            "Évaluer les limites, tiers concernés et validations requises.",
            "Préparer une réponse claire et consigner les décisions.",
        ],
    )
    doc.add_heading("4. Réponse et clôture", level=1)
    add_bullets(
        doc,
        [
            "Transmettre la réponse par un canal approprié et sécuritaire.",
            "Noter la date de réponse, le contenu remis et toute action de rectification.",
            "Mettre à jour l’inventaire des données si un écart est découvert.",
            "Conserver une preuve minimale du traitement selon la politique interne.",
        ],
    )
    doc.add_heading("Textes de formulaire à adapter", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [1.65, 5.32])
    form_rows = [
        (
            "Introduction",
            "Utilisez ce formulaire pour transmettre à {{COMPANY_NAME}} une demande concernant vos renseignements personnels.",
        ),
        (
            "Usage",
            "Les renseignements fournis servent uniquement à vérifier votre identité, traiter votre demande et assurer le suivi.",
        ),
        (
            "Champs",
            "Nom, courriel, type de demande, description, moyen de contact privilégié. Ne joignez aucun document sensible non demandé.",
        ),
        (
            "Confirmation",
            "Nous confirmerons la réception et pourrons demander une vérification raisonnable de votre identité avant de répondre.",
        ),
    ]
    for row, values in enumerate(form_rows):
        table.cell(row, 0).text = values[0]
        table.cell(row, 1).text = values[1]
        set_cell_shading(table.cell(row, 0), PALE_BLUE)
    doc.add_heading("Registre de traitement minimal", level=1)
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    set_table_widths(table, [1.1, 1.35, 1.35, 1.7, 1.47])
    for col, header in enumerate(
        ["Référence", "Reçue le", "Type", "Responsable", "État"]
    ):
        table.cell(0, col).text = header
        set_cell_shading(table.cell(0, col), BLUE)
        for run in table.cell(0, col).paragraphs[0].runs:
            set_run_font(run, color=WHITE, bold=True)
        table.cell(1, col).text = "À compléter"
    add_signoff(doc)
    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT / "procedure-demandes-loi25.docx")


def main() -> None:
    build_ai_charter()
    build_ai_memo()
    build_cyber_incident()
    build_cyber_memo()
    build_law25_requests()
    print(f"Created 5 DOCX templates in {OUTPUT}")


if __name__ == "__main__":
    main()
